"""
ActaAI - Servidor local para el asistente de reuniones
Ejecutar: python servidor.py
Luego abrir: http://localhost:8765
"""

import http.server
import socketserver
import json
import os
import base64
import tempfile
import threading
import webbrowser
import sys
import time
from pathlib import Path

PORT = 8765

# Carpeta donde se guardan las grabaciones (junto al servidor.py)
GRABACIONES_DIR = Path(__file__).parent / "grabaciones_reuniones"
GRABACIONES_DIR.mkdir(exist_ok=True)

# Intentar importar dependencias opcionales
try:
    import anthropic
    ANTHROPIC_OK = True
except ImportError:
    ANTHROPIC_OK = False
    print("⚠️  Instala anthropic: pip install anthropic")

try:
    import speech_recognition as sr
    SR_OK = True
except ImportError:
    SR_OK = False

# ──────────────────────────────────────────────
# HTML de la aplicación (embebido en el servidor)
# ──────────────────────────────────────────────
HTML = r"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>ActaAI — Asistente de Reuniones</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Sans:opsz,wght@9..40,300;9..40,400;9..40,500;9..40,600&display=swap" rel="stylesheet">
<style>
*{box-sizing:border-box;margin:0;padding:0}
:root{
  --ink:#1a1a2e;--ink2:#4a4a6a;--ink3:#9090aa;
  --paper:#f8f7f4;--surface:#eeecea;--white:#fff;
  --accent:#5b4fcf;--accent2:#e8643a;
  --green:#1e8a5e;--red:#c0392b;
  --border:#dddbd5;--shadow:0 2px 16px rgba(26,26,46,.09);
}
body{font-family:'DM Sans',sans-serif;background:var(--paper);color:var(--ink);min-height:100vh;font-size:15px}
h1,h2,h3,h4{font-family:'DM Serif Display',serif;font-weight:400}

/* ── LAYOUT ── */
.app{display:grid;grid-template-columns:240px 1fr;min-height:100vh}
.sidebar{background:var(--ink);color:#fff;padding:24px 18px;display:flex;flex-direction:column;gap:20px;position:sticky;top:0;height:100vh}
.logo{font-family:'DM Serif Display',serif;font-size:24px;color:#fff}
.logo em{color:#a89ff5;font-style:italic}
.logo-sub{font-size:11px;color:#666688;margin-top:2px}
.nav-label{font-size:10px;text-transform:uppercase;letter-spacing:2px;color:#44446a;margin-bottom:6px}
.nav-item{display:flex;align-items:center;gap:9px;padding:9px 11px;border-radius:8px;cursor:pointer;font-size:13px;color:#aaaabb;transition:all .15s;user-select:none}
.nav-item:hover{background:rgba(255,255,255,.06);color:#ddd}
.nav-item.active{background:var(--accent);color:#fff}
.nav-num{width:20px;height:20px;border-radius:50%;background:rgba(255,255,255,.1);display:flex;align-items:center;justify-content:center;font-size:11px;font-weight:600}
.nav-item.active .nav-num{background:rgba(255,255,255,.25)}
.status-pill{margin-top:auto;padding:10px 12px;background:rgba(255,255,255,.05);border-radius:10px;font-size:11px;color:#666688;line-height:1.7}
.status-pill strong{color:#888899}

/* ── MAIN ── */
.main{padding:36px 40px;max-width:860px}
.page{display:none}.page.active{display:block}
.page-title{font-size:28px;color:var(--ink);margin-bottom:6px}
.page-sub{font-size:14px;color:var(--ink3);margin-bottom:28px}

/* ── TABS ── */
.tabs{display:flex;gap:3px;background:var(--surface);border-radius:11px;padding:4px;width:fit-content;margin-bottom:24px}
.tab{padding:7px 18px;border-radius:8px;cursor:pointer;font-size:13px;font-weight:500;color:var(--ink2);transition:all .15s;border:none;background:transparent}
.tab.active{background:#fff;color:var(--ink);box-shadow:var(--shadow)}

/* ── CARDS ── */
.card{background:#fff;border:1px solid var(--border);border-radius:18px;padding:28px;margin-bottom:18px}
.card-title{font-size:13px;font-weight:600;color:var(--ink2);text-transform:uppercase;letter-spacing:.8px;margin-bottom:16px}

/* ── RECORDER ── */
.rec-center{display:flex;flex-direction:column;align-items:center;gap:18px}
.wave-wrap{width:100%;height:54px;border-radius:10px;overflow:hidden;background:var(--surface)}
canvas#wave{width:100%;height:54px;display:block}
.timer{font-family:'DM Serif Display',serif;font-size:48px;letter-spacing:3px;color:var(--ink)}
.rec-btn{width:80px;height:80px;border-radius:50%;border:none;cursor:pointer;display:flex;align-items:center;justify-content:center;transition:all .2s}
.rec-btn.idle{background:var(--accent);box-shadow:0 4px 20px rgba(91,79,207,.35)}
.rec-btn.idle:hover{transform:scale(1.06)}
.rec-btn.recording{background:var(--red);animation:pulse 1.5s infinite}
.rec-btn.paused{background:#e67e22}
@keyframes pulse{0%,100%{box-shadow:0 4px 20px rgba(192,57,43,.3)}50%{box-shadow:0 0 0 18px rgba(192,57,43,.08),0 4px 20px rgba(192,57,43,.3)}}
.rec-btn svg{width:28px;height:28px;fill:#fff}
.btn-row{display:flex;gap:10px;flex-wrap:wrap;justify-content:center}

/* ── BUTTONS ── */
.btn{padding:9px 20px;border-radius:10px;border:1px solid var(--border);background:#fff;cursor:pointer;font-size:13px;font-weight:500;color:var(--ink2);transition:all .15s;display:inline-flex;align-items:center;gap:6px;font-family:'DM Sans',sans-serif}
.btn:hover{background:var(--surface);border-color:var(--ink3)}
.btn.primary{background:var(--accent);color:#fff;border-color:var(--accent)}
.btn.primary:hover{background:#4a3fbf;border-color:#4a3fbf}
.btn.danger{border-color:#f5c6c1;color:var(--red)}
.btn.danger:hover{background:#fdf2f1}
.btn:disabled{opacity:.45;cursor:not-allowed;pointer-events:none}
.btn-full{width:100%;justify-content:center;padding:13px 20px;font-size:15px;border-radius:12px}

/* ── STATUS ── */
.status{display:flex;align-items:center;gap:9px;padding:11px 15px;border-radius:10px;font-size:13px;line-height:1.5}
.status.info{background:#eef2ff;color:#3730a3}
.status.ok{background:#ecfdf5;color:#065f46}
.status.err{background:#fef2f2;color:#991b1b}
.status.warn{background:#fffbeb;color:#92400e}
.spinner{width:15px;height:15px;border:2px solid currentColor;border-top-color:transparent;border-radius:50%;animation:spin .8s linear infinite;flex-shrink:0}
@keyframes spin{to{transform:rotate(360deg)}}

/* ── UPLOAD ── */
.drop-zone{border:2px dashed var(--border);border-radius:14px;padding:32px;text-align:center;cursor:pointer;transition:all .2s;background:var(--surface)}
.drop-zone:hover,.drop-zone.drag{border-color:var(--accent);background:#f0eeff}
.drop-icon{font-size:34px;margin-bottom:10px}
.drop-title{font-size:15px;font-weight:600;color:var(--ink);margin-bottom:3px}
.drop-sub{font-size:13px;color:var(--ink3)}

/* ── TRANSCRIPT ── */
.transcript-box{background:#fff;border:1.5px solid var(--border);border-radius:14px;padding:18px 20px;min-height:160px;font-size:14px;line-height:1.85;color:var(--ink2);outline:none;white-space:pre-wrap;transition:border-color .15s}
.transcript-box:focus{border-color:var(--accent);box-shadow:0 0 0 3px rgba(91,79,207,.08)}
.transcript-box:empty::before{content:attr(data-placeholder);color:var(--ink3);font-style:italic;pointer-events:none}

/* ── API KEY ── */
.apikey-row{display:flex;gap:8px;margin-top:10px}
.apikey-input{flex:1;border:1px solid var(--border);border-radius:9px;padding:9px 13px;font-size:13px;font-family:'DM Sans',sans-serif;color:var(--ink);outline:none;background:#fff}
.apikey-input:focus{border-color:var(--accent);box-shadow:0 0 0 3px rgba(91,79,207,.08)}

/* ── OUTPUT GRID ── */
.out-grid{display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-bottom:20px}
.out-card{background:#fff;border:1.5px solid var(--border);border-radius:16px;padding:18px;cursor:pointer;transition:all .18s}
.out-card:hover{border-color:var(--accent);transform:translateY(-2px);box-shadow:var(--shadow)}
.out-card.sel{border-color:var(--accent);box-shadow:0 0 0 3px rgba(91,79,207,.1);background:#faf9ff}
.out-emoji{font-size:26px;margin-bottom:8px}
.out-name{font-size:14px;font-weight:600;color:var(--ink);margin-bottom:3px}
.out-desc{font-size:12px;color:var(--ink3);line-height:1.5}

/* ── RESULT ── */
.result-box{background:#fff;border:1px solid var(--border);border-radius:16px;padding:24px;margin-top:18px}
.result-box h3{font-size:20px;color:var(--ink);margin-bottom:14px;display:flex;align-items:center;gap:10px}
.result-content{font-size:14px;line-height:1.9;color:var(--ink2);white-space:pre-wrap}
.result-content strong{color:var(--ink);font-weight:600}
.download-btn{margin-left:auto;padding:7px 16px;border-radius:9px;border:1px solid var(--border);background:#fff;cursor:pointer;font-size:12px;font-weight:500;color:var(--ink2);transition:all .15s;display:inline-flex;align-items:center;gap:5px}
.download-btn:hover{background:var(--accent);color:#fff;border-color:var(--accent)}

/* ── MINDMAP ── */
.mindmap-wrap{width:100%;overflow:auto;border-radius:12px;background:var(--surface);padding:8px}
svg.mm{font-family:'DM Sans',sans-serif}

@media(max-width:640px){
  .app{grid-template-columns:1fr}.sidebar{display:none}
  .out-grid{grid-template-columns:1fr}.main{padding:20px}
}
</style>
</head>
<body>
<div class="app">

<!-- SIDEBAR -->
<aside class="sidebar">
  <div>
    <div class="logo">Acta<em>AI</em></div>
    <div class="logo-sub">Asistente de reuniones</div>
  </div>
  <div>
    <div class="nav-label">Pasos</div>
    <div class="nav-item active" id="nav-0" onclick="goto(0)"><div class="nav-num">1</div> Grabar o subir</div>
    <div class="nav-item" id="nav-1" onclick="goto(1)"><div class="nav-num">2</div> Transcribir</div>
    <div class="nav-item" id="nav-2" onclick="goto(2)"><div class="nav-num">3</div> Generar outputs</div>
  </div>
  <div class="status-pill" id="sideStatus">
    <strong>Estado:</strong> Listo<br>
    <strong>Idioma:</strong> Español<br>
    <strong>Servidor:</strong> localhost:8765<br>
    <strong>API Key:</strong> <span id="keyStatus">No guardada</span>
  </div>
</aside>

<!-- MAIN -->
<main class="main">

  <!-- PAGE 0 — GRABAR -->
  <div class="page active" id="page-0">
    <h2 class="page-title">Grabar reunión</h2>
    <p class="page-sub">Graba en vivo desde el micrófono o sube un archivo de audio ya grabado.</p>

    <div class="tabs">
      <button class="tab active" onclick="switchMode('mic',this)">🎙 Micrófono en vivo</button>
      <button class="tab" onclick="switchMode('file',this)">📁 Subir archivo</button>
    </div>

    <!-- MIC -->
    <div id="mode-mic">
      <div class="card">
        <div class="rec-center">
          <div class="wave-wrap"><canvas id="wave" width="700" height="54"></canvas></div>
          <div class="timer" id="timer">00:00:00</div>
          <button class="rec-btn idle" id="recBtn" onclick="toggleRec()">
            <svg viewBox="0 0 24 24"><circle cx="12" cy="12" r="6"/></svg>
          </button>
          <div class="btn-row" id="recControls" style="display:none">
            <button class="btn" id="pauseBtn" onclick="pauseRec()">⏸ Pausar</button>
            <button class="btn danger" onclick="stopRec()">⏹ Detener y guardar</button>
          </div>
          <div id="recStatus" style="width:100%"></div>
        </div>
      </div>
    </div>

    <!-- FILE -->
    <div id="mode-file" style="display:none">
      <div class="drop-zone" id="dropZone"
           onclick="document.getElementById('fi').click()"
           ondragover="event.preventDefault();this.classList.add('drag')"
           ondragleave="this.classList.remove('drag')"
           ondrop="handleDrop(event)">
        <div class="drop-icon">🎵</div>
        <div class="drop-title">Arrastra tu archivo de audio aquí</div>
        <div class="drop-sub">o haz clic para seleccionar · MP3, WAV, M4A, OGG, WebM</div>
      </div>
      <input type="file" id="fi" accept="audio/*" style="display:none" onchange="loadFile(event)">
      <div id="fileStatus" style="margin-top:12px"></div>
    </div>

    <!-- AUDIO PREVIEW -->
    <div id="audioPreview" style="display:none;margin-top:20px">
      <div class="card">
        <div class="card-title">Audio capturado</div>
        <audio id="player" controls style="width:100%;border-radius:10px;margin-bottom:14px"></audio>
        <button class="btn primary" onclick="goto(1)">✨ Continuar a Transcribir →</button>
      </div>
    </div>

    <!-- GRABACIONES GUARDADAS -->
    <div class="card" style="margin-top:18px">
      <div class="card-title" style="display:flex;align-items:center;justify-content:space-between">
        💾 Grabaciones guardadas
        <button class="btn" style="padding:5px 12px;font-size:12px" onclick="showRecordings()">🔄 Actualizar</button>
      </div>
      <div id="recordingsList">
        <p style="color:var(--ink3);font-size:13px;padding:4px">Haz clic en "Actualizar" para ver tus grabaciones guardadas.</p>
      </div>
      <p style="font-size:12px;color:var(--ink3);margin-top:8px">Las grabaciones se guardan automáticamente en la carpeta <strong>grabaciones_reuniones/</strong> junto al servidor.py</p>
    </div>
  </div>

  <!-- PAGE 1 — TRANSCRIBIR -->
  <div class="page" id="page-1">
    <h2 class="page-title">Transcribir</h2>
    <p class="page-sub">Convierte el audio a texto. El servidor local maneja la transcripción por ti.</p>

    <!-- API KEY -->
    <div class="card" id="apiCard">
      <div class="card-title">🔑 API Key de Anthropic</div>
      <p style="font-size:13px;color:var(--ink3);margin-bottom:8px">Necesaria para generar resúmenes, mapas mentales y documentos con IA.</p>
      <div class="apikey-row">
        <input type="password" id="apiKey" class="apikey-input" placeholder="sk-ant-api03-...">
        <button class="btn primary" onclick="saveKey()">Guardar</button>
      </div>
      <div id="keyMsg" style="margin-top:10px"></div>
    </div>

    <!-- TRANSCRIPCIÓN -->
    <div class="card">
      <div class="card-title">Texto de la reunión</div>
      <div class="transcript-box" id="txBox" contenteditable="true"
           data-placeholder="La transcripción aparece aquí… también puedes pegar texto manualmente."
           oninput="onTxInput()"></div>
      <div id="txStatus" style="margin-top:12px"></div>
      <div class="btn-row" style="margin-top:14px;justify-content:flex-start">
        <button class="btn" onclick="transcribeWithSpeech()" id="speechBtn">🎤 Transcribir con voz (Chrome)</button>
        <button class="btn" onclick="transcribeFile()" id="fileBtn" style="display:none">🤖 Transcribir archivo con IA</button>
        <button class="btn danger" onclick="clearTx()">🗑 Limpiar</button>
      </div>
    </div>

    <button class="btn primary btn-full" onclick="goto(2)">✨ Generar outputs →</button>
  </div>

  <!-- PAGE 2 — GENERAR -->
  <div class="page" id="page-2">
    <h2 class="page-title">Generar outputs</h2>
    <p class="page-sub">Selecciona qué quieres generar con IA a partir de la transcripción.</p>

    <div class="out-grid">
      <div class="out-card sel" id="oc-resumen" onclick="selOut('resumen',this)">
        <div class="out-emoji">📋</div>
        <div class="out-name">Resumen ejecutivo</div>
        <div class="out-desc">Puntos clave, decisiones y contexto</div>
      </div>
      <div class="out-card" id="oc-tareas" onclick="selOut('tareas',this)">
        <div class="out-emoji">✅</div>
        <div class="out-name">Lista de tareas</div>
        <div class="out-desc">Acuerdos, responsables y fechas</div>
      </div>
      <div class="out-card" id="oc-mapa" onclick="selOut('mapa',this)">
        <div class="out-emoji">🧠</div>
        <div class="out-name">Mapa mental</div>
        <div class="out-desc">Visualización de temas y relaciones</div>
      </div>
      <div class="out-card" id="oc-acta" onclick="selOut('acta',this)">
        <div class="out-emoji">📄</div>
        <div class="out-name">Acta formal</div>
        <div class="out-desc">Documento Word listo para descargar</div>
      </div>
    </div>

    <button class="btn primary btn-full" id="genBtn" onclick="generate()">✨ Generar con IA</button>
    <div id="genStatus" style="margin-top:14px"></div>
    <div id="genResult"></div>
  </div>

</main>
</div>

<script>
// ── STATE ──
let mediaRec=null, chunks=[], recording=false, paused=false;
let timerInt=null, secs=0;
let audioBlob=null, hasFile=false;
let transcript='';
let apiKey='';
let selectedOut='resumen';
let recog=null;
let anim=null, analyser=null, dataArr=null;

// ── NAVIGATION ──
function goto(n){
  document.querySelectorAll('.page').forEach((p,i)=>p.classList.toggle('active',i===n));
  document.querySelectorAll('.nav-item').forEach((el,i)=>el.classList.toggle('active',i===n));
  if(n===1) updateFileBtn();
}

// ── MODE SWITCH ──
function switchMode(m,btn){
  document.getElementById('mode-mic').style.display=m==='mic'?'block':'none';
  document.getElementById('mode-file').style.display=m==='file'?'block':'none';
  btn.parentElement.querySelectorAll('.tab').forEach(t=>t.classList.remove('active'));
  btn.classList.add('active');
}

// ── RECORDING ──
async function toggleRec(){
  if(!recording) await startRec();
}
async function startRec(){
  try{
    const stream=await navigator.mediaDevices.getUserMedia({audio:true,video:false});
    setupViz(stream);
    chunks=[];
    mediaRec=new MediaRecorder(stream,{mimeType:MediaRecorder.isTypeSupported('audio/webm;codecs=opus')?'audio/webm;codecs=opus':'audio/webm'});
    mediaRec.ondataavailable=e=>chunks.push(e.data);
    mediaRec.onstop=()=>{
      audioBlob=new Blob(chunks,{type:'audio/webm'});
      hasFile=true;
      document.getElementById('player').src=URL.createObjectURL(audioBlob);
      document.getElementById('audioPreview').style.display='block';
      stat('recStatus','ok','⏳ Guardando grabación en tu PC...');
      saveAudioToServer(audioBlob);
    };
    mediaRec.start(250);
    recording=true;paused=false;
    document.getElementById('recBtn').className='rec-btn recording';
    document.getElementById('recBtn').innerHTML='<svg viewBox="0 0 24 24"><rect x="6" y="6" width="12" height="12" rx="2"/></svg>';
    document.getElementById('recControls').style.display='flex';
    startTimer();
  }catch(e){
    stat('recStatus','err','❌ Sin acceso al micrófono: '+e.message);
  }
}
function pauseRec(){
  if(!mediaRec) return;
  if(!paused){mediaRec.pause();paused=true;clearInterval(timerInt);document.getElementById('pauseBtn').textContent='▶ Reanudar';}
  else{mediaRec.resume();paused=false;startTimer();document.getElementById('pauseBtn').textContent='⏸ Pausar';}
  document.getElementById('recBtn').className='rec-btn '+(paused?'paused':'recording');
}
function stopRec(){
  if(mediaRec){mediaRec.stop();mediaRec.stream.getTracks().forEach(t=>t.stop());}
  recording=false;paused=false;
  clearInterval(timerInt);
  if(anim) cancelAnimationFrame(anim);
  document.getElementById('recBtn').className='rec-btn idle';
  document.getElementById('recBtn').innerHTML='<svg viewBox="0 0 24 24"><circle cx="12" cy="12" r="6"/></svg>';
  document.getElementById('recControls').style.display='none';
}
function startTimer(){
  timerInt=setInterval(()=>{
    secs++;
    const h=String(Math.floor(secs/3600)).padStart(2,'0');
    const m=String(Math.floor((secs%3600)/60)).padStart(2,'0');
    const s=String(secs%60).padStart(2,'0');
    document.getElementById('timer').textContent=`${h}:${m}:${s}`;
  },1000);
}

// ── WAVEFORM ──
function setupViz(stream){
  const ctx=new (window.AudioContext||window.webkitAudioContext)();
  analyser=ctx.createAnalyser(); analyser.fftSize=128;
  ctx.createMediaStreamSource(stream).connect(analyser);
  dataArr=new Uint8Array(analyser.frequencyBinCount);
  drawWave();
}
function drawWave(){
  anim=requestAnimationFrame(drawWave);
  const c=document.getElementById('wave'),g=c.getContext('2d');
  c.width=c.offsetWidth||700;
  g.clearRect(0,0,c.width,c.height);
  g.fillStyle='#eeecea'; g.fillRect(0,0,c.width,c.height);
  if(!analyser) return;
  analyser.getByteFrequencyData(dataArr);
  const bw=c.width/dataArr.length;
  dataArr.forEach((v,i)=>{
    const h=(v/255)*c.height;
    g.fillStyle=`hsl(${245+i*1.2},65%,62%)`;
    g.fillRect(i*bw,c.height-h,bw-1,h);
  });
}

// ── FILE UPLOAD ──
function handleDrop(ev){
  ev.preventDefault(); document.getElementById('dropZone').classList.remove('drag');
  const f=ev.dataTransfer.files[0];
  if(f&&f.type.startsWith('audio/')) setAudioFile(f);
}
function loadFile(ev){
  const f=ev.target.files[0]; if(f) setAudioFile(f);
}
function setAudioFile(f){
  audioBlob=f; hasFile=true;
  document.getElementById('player').src=URL.createObjectURL(f);
  document.getElementById('audioPreview').style.display='block';
  stat('fileStatus','ok',`✅ ${f.name} (${(f.size/1024/1024).toFixed(1)} MB)`);
}
function updateFileBtn(){
  document.getElementById('fileBtn').style.display=hasFile?'inline-flex':'none';
}

// ── API KEY ──
function saveKey(){
  const v=document.getElementById('apiKey').value.trim();
  if(!v.startsWith('sk-ant-')){stat('keyMsg','err','❌ Formato inválido. Debe empezar con sk-ant-');return;}
  apiKey=v;
  stat('keyMsg','ok','✅ API Key guardada en esta sesión.');
  document.getElementById('keyStatus').textContent='Guardada ✓';
  document.getElementById('apiCard').style.borderColor='#6ee7b7';
}


// ── GUARDAR AUDIO EN SERVIDOR ──
async function saveAudioToServer(blob){
  try{
    const fd=new FormData();
    fd.append('audio', blob, 'audio.webm');
    const res=await fetch('/save_audio',{method:'POST',body:fd});
    const data=await res.json();
    if(data.ok){
      stat('recStatus','ok',`✅ Grabación guardada: <strong>${data.filename}</strong> (${data.size_mb} MB)<br><small>📁 ${data.path}</small>`);
      loadRecordings();
    } else {
      stat('recStatus','warn','⚠️ No se pudo guardar automáticamente: '+data.error);
    }
  }catch(e){
    stat('recStatus','warn','⚠️ Servidor no disponible para guardar el audio.');
  }
}

async function loadRecordings(){
  try{
    const res=await fetch('/list_recordings',{method:'POST',body:'{}',headers:{'Content-Type':'application/json'}});
    // list_recordings espera POST vacío
  }catch(e){}
}

async function showRecordings(){
  try{
    const res=await fetch('/list_recordings',{method:'POST',headers:{'Content-Type':'application/json'},body:'{}'});
    const data=await res.json();
    const el=document.getElementById('recordingsList');
    if(!el) return;
    if(!data.files||data.files.length===0){
      el.innerHTML='<p style="color:var(--ink3);font-size:13px;padding:12px">No hay grabaciones guardadas aún.</p>';
      return;
    }
    el.innerHTML='<p style="font-size:12px;color:var(--ink3);margin-bottom:10px">📁 Carpeta: <strong>'+data.folder+'</strong></p>'+
      data.files.map(f=>`<div style="display:flex;align-items:center;gap:10px;padding:9px 12px;border-radius:9px;background:var(--surface);margin-bottom:6px;font-size:13px">
        <span>🎵</span>
        <span style="flex:1;color:var(--ink2)">${f.name}</span>
        <span style="color:var(--ink3);font-size:12px">${f.size_mb} MB</span>
      </div>`).join('');
  }catch(e){
    const el=document.getElementById('recordingsList');
    if(el) el.innerHTML='<p style="color:var(--ink3);font-size:13px;padding:12px">Inicia el servidor para ver grabaciones.</p>';
  }
}

// ── SPEECH RECOGNITION (vía servidor local HTTP) ──
function transcribeWithSpeech(){
  const SR=window.SpeechRecognition||window.webkitSpeechRecognition;
  if(!SR){
    stat('txStatus','warn','⚠️ Usa Chrome para reconocimiento de voz. También puedes pegar el texto manualmente.');
    return;
  }
  recog=new SR();
  recog.lang='es-ES'; recog.continuous=true; recog.interimResults=true; recog.maxAlternatives=1;
  let final='';
  document.getElementById('speechBtn').textContent='🔴 Escuchando... (clic para detener)';
  document.getElementById('speechBtn').onclick=()=>recog.stop();
  stat('txStatus','info','🎤 Escuchando en español... habla con claridad.');
  recog.onresult=e=>{
    let interim='';
    for(let i=e.resultIndex;i<e.results.length;i++){
      if(e.results[i].isFinal) final+=e.results[i][0].transcript+' ';
      else interim+=e.results[i][0].transcript;
    }
    setTx(final+(interim?` [${interim}]`:''));
  };
  recog.onerror=e=>{stat('txStatus','err','❌ '+e.error);resetSpeechBtn();};
  recog.onend=()=>{
    setTx(final.trim());
    stat('txStatus','ok','✅ Transcripción completada. Revisa y edita si es necesario.');
    resetSpeechBtn();
  };
  recog.start();
}
function resetSpeechBtn(){
  const b=document.getElementById('speechBtn');
  b.textContent='🎤 Transcribir con voz (Chrome)'; b.onclick=transcribeWithSpeech;
}

// ── TRANSCRIBE FILE VIA SERVER ──
async function transcribeFile(){
  if(!audioBlob){stat('txStatus','err','❌ No hay archivo de audio cargado.');return;}
  if(!apiKey){stat('txStatus','err','❌ Ingresa tu API Key primero.');return;}
  stat('txStatus','info','<div class="spinner"></div> Enviando al servidor para transcribir...');
  try{
    const fd=new FormData();
    fd.append('audio',audioBlob,audioBlob.name||'audio.webm');
    fd.append('api_key',apiKey);
    const res=await fetch('/transcribe',{method:'POST',body:fd});
    const data=await res.json();
    if(data.error) throw new Error(data.error);
    setTx(data.transcript);
    stat('txStatus','ok','✅ Transcripción completada.');
  }catch(e){
    stat('txStatus','err','❌ '+e.message);
  }
}

function setTx(t){
  const b=document.getElementById('txBox');
  b.textContent=t; transcript=t;
}
function onTxInput(){
  transcript=document.getElementById('txBox').textContent;
}
function clearTx(){
  document.getElementById('txBox').textContent=''; transcript='';
}

// ── GENERATE ──
function selOut(type,card){
  selectedOut=type;
  document.querySelectorAll('.out-card').forEach(c=>c.classList.remove('sel'));
  card.classList.add('sel');
}

async function generate(){
  transcript=document.getElementById('txBox').textContent.trim();
  if(!transcript){stat('genStatus','err','❌ Primero agrega la transcripción en el paso 2.');return;}
  if(!apiKey){stat('genStatus','err','❌ Ingresa tu API Key en el paso 2.');return;}

  const btn=document.getElementById('genBtn');
  btn.disabled=true; btn.innerHTML='<div class="spinner"></div> Generando con IA...';
  document.getElementById('genResult').innerHTML='';

  stat('genStatus','info','<div class="spinner"></div> Procesando con Claude AI...');

  try{
    const res=await fetch('/generate',{
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body:JSON.stringify({api_key:apiKey, type:selectedOut, transcript})
    });
    const data=await res.json();
    if(data.error) throw new Error(data.error);

    document.getElementById('genStatus').innerHTML='';

    if(selectedOut==='mapa') renderMindMap(data.result);
    else if(selectedOut==='acta') renderActa(data.result, data.docx_b64);
    else renderText(data.result, selectedOut);

  }catch(e){
    stat('genStatus','err','❌ '+e.message);
  }finally{
    btn.disabled=false; btn.innerHTML='✨ Generar con IA';
  }
}

// ── RENDER TEXT ──
function renderText(content, type){
  const meta={
    resumen:{icon:'📋',title:'Resumen Ejecutivo'},
    tareas:{icon:'✅',title:'Lista de Tareas y Acuerdos'},
  };
  const m=meta[type]||{icon:'📄',title:'Output'};
  const html=`<div class="result-box">
    <h3>${m.icon} ${m.title}
      <button class="download-btn" onclick="dlText(this,'${type}')">⬇ Descargar .txt</button>
    </h3>
    <div class="result-content">${fmtText(content)}</div>
  </div>`;
  document.getElementById('genResult').innerHTML=html;
  window._txt=content;
}

// ── RENDER ACTA ──
function renderActa(content, docxB64){
  let dlBtn='<button class="download-btn" onclick="dlText(this,\'acta\')">⬇ .txt</button>';
  if(docxB64){
    dlBtn=`<button class="download-btn" onclick="dlDocx('${docxB64}')">⬇ Descargar .docx</button>`;
  }
  const html=`<div class="result-box">
    <h3>📄 Acta de Reunión ${dlBtn}</h3>
    <div class="result-content">${fmtText(content)}</div>
  </div>`;
  document.getElementById('genResult').innerHTML=html;
  window._txt=content;
}

// ── RENDER MIND MAP ──
function renderMindMap(jsonStr){
  let d;
  try{
    const clean=jsonStr.replace(/```json\n?|\n?```/g,'').trim();
    d=JSON.parse(clean);
  }catch(e){
    stat('genStatus','err','❌ Error al renderizar el mapa mental. Intenta de nuevo.');
    return;
  }
  const colors=['#5b4fcf','#e8643a','#1e8a5e','#c2410c','#0369a1','#7c3aed'];
  const W=720,H=520,cx=W/2,cy=H/2;
  const temas=d.temas||[];
  const n=temas.length||1;
  let s=`<svg class="mm" viewBox="0 0 ${W} ${H}" xmlns="http://www.w3.org/2000/svg" style="width:100%;max-height:500px">
  <rect width="${W}" height="${H}" fill="#f8f7f4" rx="14"/>`;
  // central
  const ct=d.central||'Reunión';
  const ctw=Math.min(ct.length*8+24,180);
  s+=`<rect x="${cx-ctw/2}" y="${cy-24}" width="${ctw}" height="48" rx="24" fill="#1a1a2e"/>
  <text x="${cx}" y="${cy}" text-anchor="middle" dominant-baseline="middle" fill="white" font-size="13" font-weight="600">${xs(ct)}</text>`;
  temas.forEach((t,i)=>{
    const ang=(2*Math.PI*i/n)-(Math.PI/2);
    const r=170; const tx=cx+r*Math.cos(ang),ty=cy+r*Math.sin(ang);
    const col=colors[i%colors.length];
    s+=`<line x1="${cx}" y1="${cy}" x2="${tx}" y2="${ty}" stroke="${col}" stroke-width="2.5" stroke-opacity=".35"/>`;
    s+=`<ellipse cx="${tx}" cy="${ty}" rx="62" ry="26" fill="${col}" fill-opacity=".12" stroke="${col}" stroke-width="1.5"/>`;
    const tn=t.nombre||'';
    s+=`<text x="${tx}" y="${ty}" text-anchor="middle" dominant-baseline="middle" fill="${col}" font-size="11.5" font-weight="600">${xs(tn.length>15?tn.slice(0,14)+'…':tn)}</text>`;
    (t.subtemas||[]).slice(0,4).forEach((sub,j)=>{
      const sa=ang+(-0.55+j*0.38)*0.85;
      const sr=285;
      const sx=cx+sr*Math.cos(sa),sy=cy+sr*Math.sin(sa);
      s+=`<line x1="${tx}" y1="${ty}" x2="${sx}" y2="${sy}" stroke="${col}" stroke-width="1" stroke-opacity=".22" stroke-dasharray="4,3"/>`;
      s+=`<rect x="${sx-42}" y="${sy-14}" width="84" height="28" rx="14" fill="${col}" fill-opacity=".08" stroke="${col}" stroke-width="1" stroke-opacity=".3"/>`;
      const st=sub||'';
      s+=`<text x="${sx}" y="${sy}" text-anchor="middle" dominant-baseline="middle" fill="${col}" font-size="10">${xs(st.length>13?st.slice(0,12)+'…':st)}</text>`;
    });
  });
  s+='</svg>';
  document.getElementById('genResult').innerHTML=`<div class="result-box">
    <h3>🧠 Mapa Mental
      <button class="download-btn" onclick="dlSVG()">⬇ Descargar SVG</button>
    </h3>
    <div class="mindmap-wrap">${s}</div>
  </div>`;
  window._svg=s;
}

// ── DOWNLOADS ──
function dlText(btn, type){
  const names={resumen:'resumen_ejecutivo',tareas:'lista_tareas',acta:'acta_reunion'};
  const b=new Blob([window._txt||''],{type:'text/plain;charset=utf-8'});
  const a=document.createElement('a');
  a.href=URL.createObjectURL(b);
  a.download=`${names[type]||type}_${today()}.txt`;
  a.click();
}
function dlDocx(b64){
  const bin=atob(b64);
  const arr=new Uint8Array(bin.length);
  for(let i=0;i<bin.length;i++) arr[i]=bin.charCodeAt(i);
  const b=new Blob([arr],{type:'application/vnd.openxmlformats-officedocument.wordprocessingml.document'});
  const a=document.createElement('a');
  a.href=URL.createObjectURL(b);
  a.download=`acta_reunion_${today()}.docx`;
  a.click();
}
function dlSVG(){
  const b=new Blob([window._svg||''],{type:'image/svg+xml'});
  const a=document.createElement('a');
  a.href=URL.createObjectURL(b);
  a.download=`mapa_mental_${today()}.svg`;
  a.click();
}
function today(){return new Date().toISOString().slice(0,10)}

// ── UTILS ──
function stat(id,type,msg){
  const el=document.getElementById(id);
  if(!el) return;
  el.innerHTML=`<div class="status ${type}">${msg}</div>`;
}
function fmtText(t){
  return t.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;')
    .replace(/\*\*(.*?)\*\*/g,'<strong>$1</strong>')
    .replace(/\n/g,'<br>');
}
function xs(t){return t.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;')}
</script>
</body>
</html>"""

# ──────────────────────────────────────────────
# REQUEST HANDLER
# ──────────────────────────────────────────────
class Handler(http.server.BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        pass  # Silenciar logs HTTP

    def send_json(self, data, status=200):
        body = json.dumps(data, ensure_ascii=False).encode('utf-8')
        self.send_response(status)
        self.send_header('Content-Type', 'application/json; charset=utf-8')
        self.send_header('Content-Length', str(len(body)))
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self):
        if self.path in ('/', '/index.html'):
            body = HTML.encode('utf-8')
            self.send_response(200)
            self.send_header('Content-Type', 'text/html; charset=utf-8')
            self.send_header('Content-Length', str(len(body)))
            self.end_headers()
            self.wfile.write(body)
        else:
            self.send_response(404)
            self.end_headers()

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, GET, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()

    def do_POST(self):
        if self.path == '/generate':
            self.handle_generate()
        elif self.path == '/transcribe':
            self.handle_transcribe()
        elif self.path == '/save_audio':
            self.handle_save_audio()
        elif self.path == '/list_recordings':
            self.handle_list_recordings()
        else:
            self.send_json({'error': 'Ruta no encontrada'}, 404)

    def read_json_body(self):
        length = int(self.headers.get('Content-Length', 0))
        return json.loads(self.rfile.read(length).decode('utf-8'))

    def handle_generate(self):
        if not ANTHROPIC_OK:
            self.send_json({'error': 'Instala anthropic: pip install anthropic'})
            return
        try:
            data = self.read_json_body()
            api_key = data.get('api_key', '')
            out_type = data.get('type', 'resumen')
            transcript = data.get('transcript', '')

            if not api_key:
                self.send_json({'error': 'API Key requerida'})
                return
            if not transcript:
                self.send_json({'error': 'Transcripción vacía'})
                return

            prompts = {
                'resumen': f"""Eres un asistente experto en gestión de reuniones. Analiza esta transcripción en español y genera un RESUMEN EJECUTIVO con las siguientes secciones bien formateadas:

**1. Resumen general**
(2-3 oraciones que describan el propósito y resultado de la reunión)

**2. Temas principales tratados**
(lista de los temas clave discutidos)

**3. Decisiones tomadas**
(lista de decisiones concretas adoptadas)

**4. Compromisos y acuerdos**
(lista de compromisos asumidos)

**5. Próximos pasos**
(acciones inmediatas a realizar)

Transcripción:
{transcript}""",

                'tareas': f"""Eres un asistente de gestión de proyectos. Analiza esta transcripción de reunión en español y extrae TODAS las tareas, compromisos y acuerdos mencionados.

Formatea la respuesta así:

**TAREAS Y COMPROMISOS**

Para cada tarea usa este formato:
• [Tarea]: descripción clara
  - Responsable: nombre o "Por definir"
  - Fecha límite: fecha o "Por definir"
  - Prioridad: Alta / Media / Baja

Al final incluye una sección **ACUERDOS GENERALES** con los puntos acordados por el grupo.

Transcripción:
{transcript}""",

                'mapa': f"""Analiza esta transcripción de reunión en español y extrae los temas principales con sus subtemas.

Devuelve ÚNICAMENTE un JSON válido con este formato exacto, sin texto adicional ni backticks:
{{
  "central": "Título breve de la reunión (máx 4 palabras)",
  "temas": [
    {{"nombre": "Tema principal 1 (máx 4 palabras)", "subtemas": ["Subtema 1a (máx 4 palabras)", "Subtema 1b"]}},
    {{"nombre": "Tema principal 2", "subtemas": ["Subtema 2a", "Subtema 2b"]}},
    {{"nombre": "Tema principal 3", "subtemas": ["Subtema 3a"]}}
  ]
}}

Reglas: máximo 5 temas principales, máximo 4 subtemas por tema, textos MUY cortos.

Transcripción:
{transcript}""",

                'acta': f"""Eres un asistente administrativo experto. Genera un ACTA DE REUNIÓN formal y profesional en español a partir de esta transcripción.

Estructura requerida:

**ACTA DE REUNIÓN**
Fecha: {__import__('datetime').date.today().strftime('%d de %B de %Y')}
Modalidad: [presencial/virtual según contexto]
Participantes: [lista los mencionados o "Ver asistentes"]

**1. APERTURA**
[Breve descripción del inicio]

**2. ORDEN DEL DÍA**
[Lista los temas tratados]

**3. DESARROLLO**
[Descripción detallada por tema]

**4. ACUERDOS Y COMPROMISOS**
[Lista numerada de acuerdos]

**5. CIERRE**
[Conclusión y próxima reunión si se menciona]

___________________________
Firma y sello

Transcripción:
{transcript}"""
            }

            client = anthropic.Anthropic(api_key=api_key)
            message = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=2500,
                messages=[{"role": "user", "content": prompts.get(out_type, prompts['resumen'])}]
            )
            result = message.content[0].text

            response = {'result': result}

            # Generar .docx si es acta
            if out_type == 'acta':
                docx_b64 = generate_docx(result, transcript)
                if docx_b64:
                    response['docx_b64'] = docx_b64

            self.send_json(response)

        except anthropic.AuthenticationError:
            self.send_json({'error': 'API Key inválida. Verifica que sea correcta.'})
        except Exception as e:
            self.send_json({'error': str(e)})

    def handle_transcribe(self):
        """Transcribir archivo de audio usando Whisper vía Claude o speech_recognition"""
        try:
            import cgi
            ctype = self.headers.get('Content-Type', '')
            length = int(self.headers.get('Content-Length', 0))
            raw = self.rfile.read(length)

            # Parsear multipart
            import email
            from io import BytesIO
            msg_str = f'Content-Type: {ctype}\r\n\r\n'.encode() + raw
            msg = email.message_from_bytes(msg_str)

            audio_data = None
            api_key_val = ''

            for part in msg.walk():
                name = part.get_param('name', header='content-disposition')
                if name == 'audio':
                    audio_data = part.get_payload(decode=True)
                elif name == 'api_key':
                    api_key_val = part.get_payload(decode=True).decode('utf-8').strip()

            if not audio_data:
                self.send_json({'error': 'No se recibió audio'})
                return

            # Guardar audio temporal y transcribir con speech_recognition
            with tempfile.NamedTemporaryFile(suffix='.webm', delete=False) as f:
                f.write(audio_data)
                tmp_path = f.name

            transcript_text = ''

            # Intento con speech_recognition
            if SR_OK:
                try:
                    rec = sr.Recognizer()
                    # Convertir a wav si es necesario
                    wav_path = tmp_path.replace('.webm', '.wav')
                    os.system(f'ffmpeg -i {tmp_path} {wav_path} -y -loglevel quiet 2>/dev/null')
                    use_path = wav_path if os.path.exists(wav_path) else tmp_path
                    with sr.AudioFile(use_path) as source:
                        audio = rec.record(source)
                    transcript_text = rec.recognize_google(audio, language='es-ES')
                except Exception as e:
                    transcript_text = ''

            # Si no funcionó, pedir al usuario que use el método manual
            if not transcript_text:
                self.send_json({
                    'error': 'Para transcribir archivos instala: pip install SpeechRecognition ffmpeg-python\n'
                             'O usa el botón del micrófono para grabar en vivo.'
                })
                return

            os.unlink(tmp_path)
            self.send_json({'transcript': transcript_text})

        except Exception as e:
            self.send_json({'error': f'Error al transcribir: {str(e)}'})


    def handle_save_audio(self):
        """Guarda el audio grabado en la carpeta grabaciones_reuniones/"""
        try:
            import datetime, email
            ctype = self.headers.get('Content-Type', '')
            length = int(self.headers.get('Content-Length', 0))
            raw = self.rfile.read(length)

            # Parsear multipart
            msg_str = f'Content-Type: {ctype}\r\n\r\n'.encode() + raw
            msg = email.message_from_bytes(msg_str)

            audio_data = None
            filename_hint = ''

            for part in msg.walk():
                name = part.get_param('name', header='content-disposition')
                if name == 'audio':
                    audio_data = part.get_payload(decode=True)
                elif name == 'filename':
                    filename_hint = part.get_payload(decode=True).decode('utf-8').strip()

            if not audio_data:
                self.send_json({'error': 'No se recibió audio'})
                return

            ts = datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
            ext = '.webm'
            if filename_hint:
                ext = Path(filename_hint).suffix or ext
            fname = f'reunion_{ts}{ext}'
            fpath = GRABACIONES_DIR / fname
            fpath.write_bytes(audio_data)

            size_mb = len(audio_data) / 1024 / 1024
            print(f'  💾  Audio guardado: {fname} ({size_mb:.1f} MB)')
            self.send_json({
                'ok': True,
                'filename': fname,
                'path': str(fpath),
                'size_mb': round(size_mb, 2)
            })

        except Exception as e:
            self.send_json({'error': f'Error al guardar: {str(e)}'})

    def handle_list_recordings(self):
        """Devuelve la lista de grabaciones guardadas"""
        try:
            files = []
            for f in sorted(GRABACIONES_DIR.glob('*'), reverse=True):
                if f.is_file():
                    files.append({
                        'name': f.name,
                        'size_mb': round(f.stat().st_size / 1024 / 1024, 2),
                        'date': f.name.replace('reunion_', '').replace('_', ' ').split('.')[0]
                    })
            self.send_json({'files': files[:20], 'folder': str(GRABACIONES_DIR)})
        except Exception as e:
            self.send_json({'error': str(e)})


def generate_docx(content, transcript=''):
    """Genera un .docx en memoria y lo devuelve en base64"""
    try:
        import subprocess, sys, importlib

        # Instalar docx si no está
        try:
            import docx as _d
        except ImportError:
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])

        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io, datetime

        doc = Document()

        # Estilos
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Título
        title = doc.add_heading('ACTA DE REUNIÓN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # Fecha
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Fecha: {datetime.date.today().strftime("%d/%m/%Y")}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x88)

        doc.add_paragraph()

        # Contenido generado por IA
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith('**') and line.endswith('**'):
                h = doc.add_heading(line.strip('*'), level=2)
                h.runs[0].font.color.rgb = RGBColor(0x5b, 0x4f, 0xcf)
            elif line.startswith('• ') or line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])
            elif line[0].isdigit() and '. ' in line[:4]:
                p = doc.add_paragraph(style='List Number')
                p.add_run(line[line.index('. ')+2:])
            else:
                # Procesar **bold**
                p = doc.add_paragraph()
                import re
                parts = re.split(r'\*\*(.*?)\*\*', line)
                for k, part in enumerate(parts):
                    run = p.add_run(part)
                    if k % 2 == 1:
                        run.bold = True

        # Guardar en buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return base64.b64encode(buf.read()).decode('utf-8')

    except Exception as e:
        print(f'  ⚠️  No se pudo generar .docx: {e}')
        return None


# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────
def main():
    print("\n" + "="*50)
    print("  🎙  ActaAI — Asistente de Reuniones")
    print("="*50)

    if not ANTHROPIC_OK:
        print("\n  ⚠️  Instala dependencias primero:")
        print("     pip install anthropic python-docx\n")

    print(f"\n  ✅  Servidor iniciado en http://localhost:{PORT}")
    print("  📌  Abre esa URL en Chrome para usar la app")
    print(f"  💾  Grabaciones en: {GRABACIONES_DIR}")
    print("  🛑  Presiona Ctrl+C para detener\n")

    # Abrir navegador automáticamente tras 1 segundo
    def open_browser():
        time.sleep(1.2)
        webbrowser.open(f'http://localhost:{PORT}')
    threading.Thread(target=open_browser, daemon=True).start()

    with socketserver.TCPServer(("", PORT), Handler) as httpd:
        httpd.allow_reuse_address = True
        try:
            httpd.serve_forever()
        except KeyboardInterrupt:
            print("\n  👋  Servidor detenido.\n")

if __name__ == '__main__':
    main()
